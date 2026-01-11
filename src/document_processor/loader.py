"""Document loader supporting multiple file formats."""

import os
from pathlib import Path
from typing import List, Optional
from dataclasses import dataclass, field


@dataclass
class Document:
    """Represents a loaded document with metadata."""
    
    content: str
    metadata: dict = field(default_factory=dict)
    
    @property
    def source(self) -> str:
        return self.metadata.get("source", "unknown")
    
    @property
    def page(self) -> Optional[int]:
        return self.metadata.get("page")


class DocumentLoader:
    """
    Load documents from various sources (PDF, DOCX, TXT, web URLs).
    
    Supports:
    - PDF files (via pypdf)
    - Word documents (via python-docx)
    - Text files
    - Web URLs (via requests + beautifulsoup)
    - Azure Document Intelligence (optional, for complex docs)
    """
    
    def __init__(
        self,
        use_azure_doc_intelligence: bool = False,
        azure_endpoint: Optional[str] = None,
        azure_key: Optional[str] = None,
    ):
        """
        Initialize the document loader.
        
        Args:
            use_azure_doc_intelligence: Use Azure Document Intelligence for extraction
            azure_endpoint: Azure Document Intelligence endpoint
            azure_key: Azure Document Intelligence API key
        """
        self.use_azure_doc_intelligence = use_azure_doc_intelligence
        self.azure_endpoint = azure_endpoint
        self.azure_key = azure_key
        
        if use_azure_doc_intelligence and (not azure_endpoint or not azure_key):
            raise ValueError(
                "Azure Document Intelligence requires endpoint and key"
            )
    
    def load(self, source: str) -> List[Document]:
        """
        Load document(s) from a source path or URL.
        
        Args:
            source: File path, directory path, or URL
            
        Returns:
            List of Document objects
        """
        if source.startswith(("http://", "https://")):
            return self._load_url(source)
        
        path = Path(source)
        
        if path.is_dir():
            return self._load_directory(path)
        elif path.is_file():
            return self._load_file(path)
        else:
            raise FileNotFoundError(f"Source not found: {source}")
    
    def _load_directory(self, directory: Path) -> List[Document]:
        """Load all supported documents from a directory."""
        documents = []
        supported_extensions = {".pdf", ".docx", ".doc", ".txt", ".md"}
        
        for file_path in directory.rglob("*"):
            if file_path.suffix.lower() in supported_extensions:
                try:
                    documents.extend(self._load_file(file_path))
                except Exception as e:
                    print(f"Warning: Failed to load {file_path}: {e}")
        
        return documents
    
    def _load_file(self, file_path: Path) -> List[Document]:
        """Load a single file based on its extension."""
        extension = file_path.suffix.lower()
        
        if self.use_azure_doc_intelligence:
            return self._load_with_azure_doc_intelligence(file_path)
        
        if extension == ".pdf":
            return self._load_pdf(file_path)
        elif extension in {".docx", ".doc"}:
            return self._load_docx(file_path)
        elif extension in {".txt", ".md"}:
            return self._load_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def _load_pdf(self, file_path: Path) -> List[Document]:
        """Load PDF using pypdf."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf is required for PDF loading. Install with: pip install pypdf")
        
        documents = []
        reader = PdfReader(str(file_path))
        
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text.strip():
                documents.append(Document(
                    content=text,
                    metadata={
                        "source": str(file_path),
                        "page": page_num,
                        "total_pages": len(reader.pages),
                        "file_type": "pdf"
                    }
                ))
        
        return documents
    
    def _load_docx(self, file_path: Path) -> List[Document]:
        """Load Word document using python-docx."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required for DOCX loading. Install with: pip install python-docx")
        
        doc = DocxDocument(str(file_path))
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        return [Document(
            content=full_text,
            metadata={
                "source": str(file_path),
                "file_type": "docx"
            }
        )]
    
    def _load_text(self, file_path: Path) -> List[Document]:
        """Load plain text or markdown file."""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        return [Document(
            content=content,
            metadata={
                "source": str(file_path),
                "file_type": file_path.suffix.lstrip(".")
            }
        )]
    
    def _load_url(self, url: str) -> List[Document]:
        """Load content from a web URL."""
        try:
            import requests
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError(
                "requests and beautifulsoup4 are required for URL loading. "
                "Install with: pip install requests beautifulsoup4"
            )
        
        response = requests.get(url, timeout=30)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, "html.parser")
        
        # Remove script and style elements
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()
        
        text = soup.get_text(separator="\n", strip=True)
        
        return [Document(
            content=text,
            metadata={
                "source": url,
                "file_type": "html"
            }
        )]
    
    def _load_with_azure_doc_intelligence(self, file_path: Path) -> List[Document]:
        """Load document using Azure Document Intelligence."""
        try:
            from azure.ai.formrecognizer import DocumentAnalysisClient
            from azure.core.credentials import AzureKeyCredential
        except ImportError:
            raise ImportError(
                "azure-ai-formrecognizer is required. "
                "Install with: pip install azure-ai-formrecognizer"
            )
        
        client = DocumentAnalysisClient(
            endpoint=self.azure_endpoint,
            credential=AzureKeyCredential(self.azure_key)
        )
        
        with open(file_path, "rb") as f:
            poller = client.begin_analyze_document("prebuilt-read", f)
        
        result = poller.result()
        
        documents = []
        for page_num, page in enumerate(result.pages, start=1):
            page_text = ""
            for line in page.lines:
                page_text += line.content + "\n"
            
            if page_text.strip():
                documents.append(Document(
                    content=page_text,
                    metadata={
                        "source": str(file_path),
                        "page": page_num,
                        "total_pages": len(result.pages),
                        "file_type": file_path.suffix.lstrip("."),
                        "extraction_method": "azure_doc_intelligence"
                    }
                ))
        
        return documents
